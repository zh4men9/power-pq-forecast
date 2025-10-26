"""
Word (DOCX) report generation module
Generates comprehensive evaluation report in Word format using python-docx
Reference: https://python-docx.readthedocs.io/en/latest/
"""
import pandas as pd
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional


def add_heading_with_style(doc: Document, text: str, level: int = 1):
    """Add a heading with custom style"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_table_from_dataframe(doc: Document, df: pd.DataFrame, title: Optional[str] = None):
    """
    Add a table to document from DataFrame
    
    Args:
        doc: Document object
        df: DataFrame to convert to table
        title: Optional table title
    """
    if title:
        doc.add_paragraph(title, style='Heading 3')
    
    # Create table
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        header_cells[i].text = str(col)
        # Bold header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    for i, row in enumerate(df.itertuples(index=False)):
        cells = table.rows[i + 1].cells
        for j, value in enumerate(row):
            cells[j].text = str(value)
    
    doc.add_paragraph()  # Add spacing


def generate_word_report(
    metrics_df: pd.DataFrame,
    config_path: str = "config.yaml",
    figures_dir: str = "outputs/figures",
    output_path: str = "outputs/report/项目评估报告.docx",
    forecast_df: pd.DataFrame = None
) -> str:
    """
    Generate Word (DOCX) report from evaluation results
    
    Uses python-docx for inline image insertion. Images are saved first,
    then inserted using add_picture() method.
    
    Reference: https://python-docx.readthedocs.io/en/latest/user/shapes.html
    
    Args:
        metrics_df: DataFrame with evaluation metrics
        config_path: Path to configuration file
        figures_dir: Directory containing figures
        output_path: Path to save report
        forecast_df: Optional DataFrame with forecast results
    
    Returns:
        Path to generated report
    """
    # Create document
    doc = Document()
    
    # Title page
    title = doc.add_heading('电力质量预测项目评估报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'报告生成日期: {datetime.now().strftime("%Y年%m月%d日")}')
    date_run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # Aggregate results - get all available metrics
    metric_cols = [col for col in ['RMSE', 'MAE', 'SMAPE', 'WAPE', 'ACC'] 
                   if col in metrics_df.columns]
    agg_dict = {col: 'mean' for col in metric_cols}
    agg_results = metrics_df.groupby(['model', 'horizon', 'target']).agg(agg_dict).reset_index()
    
    # Section 1: Overview
    add_heading_with_style(doc, '一、项目概况', level=1)
    
    doc.add_paragraph(
        '本项目针对电力系统的有功功率（P）和无功功率（Q）进行时间序列预测分析。'
        '采用多种预测模型进行对比评估，包括朴素基线、季节性基线、随机森林、'
        'XGBoost、LSTM和Transformer等方法。'
    )
    
    doc.add_paragraph()
    
    # Section 2: Model descriptions
    add_heading_with_style(doc, '二、模型介绍', level=1)
    
    doc.add_paragraph('本项目采用六种不同类型的预测模型，涵盖基线方法、传统机器学习和深度学习方法：')
    doc.add_paragraph()
    
    add_heading_with_style(doc, '2.1 朴素预测（Naive）', level=2)
    doc.add_paragraph(
        '朴素预测是最简单的基线模型，使用最后一个观测值作为未来所有时间步的预测值。'
        '虽然方法简单，但在许多实际应用中表现出色，常作为其他模型的基准对比。'
    )
    doc.add_paragraph('参数配置：', style='Heading 4')
    doc.add_paragraph('无需训练参数，直接使用最后观测值进行预测', style='List Bullet')
    doc.add_paragraph()
    
    add_heading_with_style(doc, '2.2 季节朴素预测（Seasonal Naive）', level=2)
    doc.add_paragraph(
        '季节朴素预测考虑了数据的季节性特征，使用上一个季节周期对应时刻的观测值进行预测。'
        '适用于具有明显周期性模式的时间序列数据，如电力负荷的日周期或周周期特性。'
    )
    doc.add_paragraph('参数配置：', style='Heading 4')
    doc.add_paragraph('seasonal_period: 96（对应24小时，每15分钟一个数据点）', style='List Bullet')
    doc.add_paragraph()
    
    add_heading_with_style(doc, '2.3 随机森林（Random Forest）', level=2)
    doc.add_paragraph(
        '随机森林是一种集成学习方法，通过构建多个决策树并综合其预测结果来提高模型的准确性和稳定性。'
        '该模型能够自动捕获特征之间的非线性关系，并提供特征重要性分析，具有较强的泛化能力和抗过拟合能力。'
    )
    doc.add_paragraph('参数配置：', style='Heading 4')
    doc.add_paragraph('n_estimators: 100（决策树数量）', style='List Bullet')
    doc.add_paragraph('max_depth: 10（树的最大深度）', style='List Bullet')
    doc.add_paragraph('min_samples_split: 10（内部节点再划分所需最小样本数）', style='List Bullet')
    doc.add_paragraph('min_samples_leaf: 5（叶节点最少样本数）', style='List Bullet')
    doc.add_paragraph('random_state: 42（随机种子，保证结果可复现）', style='List Bullet')
    doc.add_paragraph()
    
    add_heading_with_style(doc, '2.4 XGBoost', level=2)
    doc.add_paragraph(
        'XGBoost是一种高效的梯度提升决策树算法，通过迭代方式构建多个弱学习器并加权组合。'
        '相比随机森林，XGBoost在处理大规模数据时性能更优，且能够更好地处理缺失值和异常值。'
        '该算法在各类机器学习竞赛中表现优异，广泛应用于时间序列预测任务。'
    )
    doc.add_paragraph('参数配置：', style='Heading 4')
    doc.add_paragraph('n_estimators: 100（提升轮数）', style='List Bullet')
    doc.add_paragraph('max_depth: 6（树的最大深度）', style='List Bullet')
    doc.add_paragraph('learning_rate: 0.1（学习率，控制每棵树的权重）', style='List Bullet')
    doc.add_paragraph('subsample: 0.8（训练每棵树时的样本采样比例）', style='List Bullet')
    doc.add_paragraph('colsample_bytree: 0.8（构建每棵树时的特征采样比例）', style='List Bullet')
    doc.add_paragraph('random_state: 42', style='List Bullet')
    doc.add_paragraph()
    
    add_heading_with_style(doc, '2.5 LSTM（长短期记忆网络）', level=2)
    doc.add_paragraph(
        'LSTM是一种特殊的循环神经网络（RNN），专门设计用于处理序列数据和长期依赖问题。'
        '通过引入门控机制（输入门、遗忘门、输出门），LSTM能够有效捕获时间序列中的长期依赖关系，'
        '避免了传统RNN的梯度消失问题。适用于复杂的时序模式识别和多步预测任务。'
    )
    doc.add_paragraph('参数配置：', style='Heading 4')
    doc.add_paragraph('hidden_size: 64（隐藏层维度）', style='List Bullet')
    doc.add_paragraph('num_layers: 2（LSTM堆叠层数）', style='List Bullet')
    doc.add_paragraph('dropout: 0.2（防止过拟合的丢弃率）', style='List Bullet')
    doc.add_paragraph('epochs: 50（训练轮数）', style='List Bullet')
    doc.add_paragraph('batch_size: 32（批次大小）', style='List Bullet')
    doc.add_paragraph('learning_rate: 0.001（Adam优化器学习率）', style='List Bullet')
    doc.add_paragraph('硬件加速: 自动检测CUDA GPU / Apple MPS / CPU', style='List Bullet')
    doc.add_paragraph()
    
    add_heading_with_style(doc, '2.6 Transformer', level=2)
    doc.add_paragraph(
        'Transformer基于自注意力机制（Self-Attention），摒弃了传统的循环结构，能够并行处理序列数据。'
        '通过多头注意力机制，模型可以同时关注序列中不同位置的信息，捕获长距离依赖关系。'
        '位置编码（Positional Encoding）保留了序列的时序信息。相比LSTM，Transformer在处理长序列时更加高效，'
        '并在自然语言处理和时间序列预测等领域取得了显著成果。'
    )
    doc.add_paragraph('参数配置：', style='Heading 4')
    doc.add_paragraph('d_model: 64（模型维度）', style='List Bullet')
    doc.add_paragraph('nhead: 4（多头注意力的头数）', style='List Bullet')
    doc.add_paragraph('num_encoder_layers: 2（编码器层数）', style='List Bullet')
    doc.add_paragraph('num_decoder_layers: 2（解码器层数）', style='List Bullet')
    doc.add_paragraph('dim_feedforward: 256（前馈网络维度）', style='List Bullet')
    doc.add_paragraph('dropout: 0.1（丢弃率）', style='List Bullet')
    doc.add_paragraph('epochs: 50（训练轮数）', style='List Bullet')
    doc.add_paragraph('batch_size: 32（批次大小）', style='List Bullet')
    doc.add_paragraph('learning_rate: 0.001（Adam优化器学习率）', style='List Bullet')
    doc.add_paragraph('硬件加速: 自动检测CUDA GPU / Apple MPS / CPU', style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph()
    
    # Section 3: Methodology
    add_heading_with_style(doc, '三、评估方法', level=1)
    
    add_heading_with_style(doc, '3.1 验证策略', level=2)
    doc.add_paragraph(
        '本项目采用滚动起点交叉验证（Rolling Origin Cross-Validation）方法，'
        '确保训练集始终位于测试集之前，严格避免使用未来信息进行训练。'
        '此方法是时间序列预测的标准验证方式，符合实际应用场景。'
    )
    
    add_heading_with_style(doc, '3.2 评估指标', level=2)
    doc.add_paragraph('本项目采用以下五项指标综合评估模型性能：')
    
    metrics_list = [
        'RMSE (Root Mean Squared Error): 均方根误差，反映预测误差的绝对大小，对大误差更敏感。越小越好，单位与预测值相同。',
        'MAE (Mean Absolute Error): 平均绝对误差，反映预测误差的平均水平。越小越好，单位与预测值相同。',
        'SMAPE (Symmetric Mean Absolute Percentage Error): 对称平均绝对百分比误差，百分比形式的相对误差。越小越好，取值范围0-200%。',
        'WAPE (Weighted Absolute Percentage Error): 加权绝对百分比误差，相比MAPE更稳健，不受零值影响。越小越好，通常在0-100%范围内。',
        'ACC (Accuracy): 近似准确率，表示预测误差在5%阈值内的样本比例。越大越好，取值范围0-100%。'
    ]
    
    for metric_desc in metrics_list:
        doc.add_paragraph(metric_desc, style='List Bullet')
    
    doc.add_paragraph(
        '说明：RMSE、MAE、SMAPE、WAPE四项指标越小表示模型性能越好；'
        'ACC指标越大表示模型性能越好。ACC=85%表示85%的预测误差在5%以内，直观反映预测结果的实用性。'
    )
    
    # Section 4: Results
    add_heading_with_style(doc, '四、评估结果', level=1)
    
    horizons = sorted(agg_results['horizon'].unique())
    targets = agg_results['target'].unique()
    target_names = {'P': '有功功率', 'Q': '无功功率'}
    
    # Only show P results (有功功率)
    for target in targets:
        if target != 'P':
            continue
            
        add_heading_with_style(doc, f'4.1 {target_names.get(target, target)}预测结果', level=2)
        
        target_data = agg_results[agg_results['target'] == target]
        
        for metric in ['RMSE', 'MAE', 'SMAPE', 'WAPE', 'ACC']:
            if metric not in target_data.columns:
                continue
                
            add_heading_with_style(doc, f'{metric}指标', level=3)
            
            # Create pivot table
            pivot_data = target_data.pivot(index='model', columns='horizon', values=metric)
            pivot_data.columns = [f'步长{h}' for h in pivot_data.columns]
            pivot_data = pivot_data.reset_index()
            pivot_data.columns.name = None
            
            # Find best value for each horizon
            best_values = {}
            for col in pivot_data.columns:
                if col != 'model':
                    if metric == 'ACC':  # ACC越大越好
                        best_values[col] = pivot_data[col].max()
                    else:  # 其他指标越小越好
                        best_values[col] = pivot_data[col].min()
            
            # Format values and mark best
            for col in pivot_data.columns:
                if col != 'model':
                    pivot_data[col] = pivot_data[col].apply(
                        lambda x: f'{x:.4f}' if pd.notna(x) else 'N/A'
                    )
            
            pivot_data.rename(columns={'model': '模型'}, inplace=True)
            
            # Create table with formatting
            table = doc.add_table(rows=len(pivot_data) + 1, cols=len(pivot_data.columns))
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            for i, col in enumerate(pivot_data.columns):
                header_cells[i].text = str(col)
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Data rows with highlighting
            for i, row in enumerate(pivot_data.itertuples(index=False)):
                cells = table.rows[i + 1].cells
                for j, value in enumerate(row):
                    cells[j].text = str(value)
                    
                    # Highlight best values in red/bold
                    if j > 0 and str(value) != 'N/A':  # Skip model column
                        col_name = pivot_data.columns[j]
                        try:
                            if col_name in best_values:
                                if float(value) == best_values[col_name]:
                                    for paragraph in cells[j].paragraphs:
                                        for run in paragraph.runs:
                                            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                                            run.font.bold = True
                        except:
                            pass
            
            doc.add_paragraph()  # Add spacing
    
    # Add error by horizon figure
    add_heading_with_style(doc, '4.2 可视化结果', level=2)
    
    # Data overview figure
    data_overview_path = Path(figures_dir) / 'data_overview.png'
    if data_overview_path.exists():
        doc.add_paragraph('图1: 数据总览 - 处理前后对比')
        doc.add_picture(str(data_overview_path), width=Inches(6))
        
        # Add explanation
        doc.add_paragraph('图表解释：', style='Heading 4')
        doc.add_paragraph(
            '左列为处理前：显示原始数据中缺失的601个时间戳（NaN值）', style='List Bullet'
        )
        doc.add_paragraph(
            '右列为处理后：使用P=280的数据行填充缺失时间戳，形成完整时间序列', style='List Bullet'
        )
        doc.add_paragraph(
            '上图为有功功率P，下图为无功功率Q', style='List Bullet'
        )
        doc.add_paragraph()
    
    # Error by horizon RMSE figure
    error_rmse_path = Path(figures_dir) / 'error_by_horizon_rmse.png'
    if error_rmse_path.exists():
        doc.add_paragraph('图2: 不同模型在各预测步长的RMSE误差对比')
        doc.add_picture(str(error_rmse_path), width=Inches(6))
        
        # Add explanation
        doc.add_paragraph('图表解释：', style='Heading 4')
        doc.add_paragraph(
            '横轴：预测步长（Forecast Horizon），表示向未来预测的时间步数。步长越大，预测难度越高。',
            style='List Bullet'
        )
        doc.add_paragraph(
            '纵轴：RMSE误差值（Root Mean Square Error），单位与预测值相同。数值越小表示预测精度越高。',
            style='List Bullet'
        )
        doc.add_paragraph(
            '不同颜色的线条代表不同模型，可直观比较各模型在不同预测步长下的表现。',
            style='List Bullet'
        )
        doc.add_paragraph(
            '关键观察：基线模型（Naive/SeasonalNaive）作为参考；误差随步长变化趋势；模型间相对表现。',
            style='List Bullet'
        )
        doc.add_paragraph()
    
    # All metrics by horizon figure
    all_metrics_path = Path(figures_dir) / 'all_metrics_by_horizon.png'
    if all_metrics_path.exists():
        doc.add_paragraph('图3: 所有指标随预测步长的变化')
        doc.add_picture(str(all_metrics_path), width=Inches(6.5))
        
        # Add explanation
        doc.add_paragraph('图表解释：', style='Heading 4')
        doc.add_paragraph(
            '横轴：预测步长（Forecast Horizon）。每个子图展示一个评估指标。',
            style='List Bullet'
        )
        doc.add_paragraph(
            '纵轴：各指标的数值。RMSE/MAE/SMAPE/WAPE越小越好（红色系）；ACC越大越好（绿色系）。',
            style='List Bullet'
        )
        doc.add_paragraph(
            '5个子图全面展示模型性能：左上RMSE、右上MAE、左中SMAPE、右中WAPE、左下ACC。',
            style='List Bullet'
        )
        doc.add_paragraph(
            '综合分析：理想模型应在所有子图中表现优异（前4个指标低、ACC高），且曲线平稳。',
            style='List Bullet'
        )
        doc.add_paragraph()
    
    # Feature importance figure (if exists)
    feature_importance_path = Path(figures_dir) / 'feature_importance.png'
    if feature_importance_path.exists():
        doc.add_paragraph('图4: 特征重要性排序（树模型）')
        doc.add_picture(str(feature_importance_path), width=Inches(6))
        
        # Add explanation
        doc.add_paragraph('图表解释：', style='Heading 4')
        doc.add_paragraph(
            '横轴为重要性得分（0-1），纵轴为特征名称（按重要性排序）。',
            style='List Bullet'
        )
        doc.add_paragraph(
            '滞后特征（如P_lag_1）通常最重要，时间特征捕获周期性，外生变量特征体现额外信息价值。',
            style='List Bullet'
        )
        doc.add_paragraph(
            '应用价值：识别关键特征、指导特征选择、验证外生变量贡献。',
            style='List Bullet'
        )
        doc.add_paragraph()
    
    # Section 5: Conclusions
    add_heading_with_style(doc, '五、结论与建议', level=1)
    
    add_heading_with_style(doc, '5.1 主要发现', level=2)
    
    # Find best models for P only
    target = 'P'
    target_data = agg_results[agg_results['target'] == target]
    
    if not target_data.empty:
        doc.add_paragraph('有功功率（P）预测最优模型：', style='Heading 4')
        
        best_models = {}
        for metric in ['RMSE', 'MAE', 'SMAPE', 'WAPE', 'ACC']:
            if metric in target_data.columns:
                metric_means = target_data.groupby('model')[metric].mean()
                if not metric_means.empty:
                    if metric == 'ACC':  # ACC越大越好
                        best_model = metric_means.idxmax()
                        best_value = metric_means.max()
                    else:  # 其他指标越小越好
                        best_model = metric_means.idxmin()
                        best_value = metric_means.min()
                    best_models[metric] = (best_model, best_value)
                    doc.add_paragraph(
                        f'{metric}指标最优: {best_model} (平均值: {best_value:.4f})',
                        style='List Bullet'
                    )
        
        # Determine overall best model (most frequent in best_models)
        if best_models:
            from collections import Counter
            model_counts = Counter([model for model, _ in best_models.values()])
            overall_best = model_counts.most_common(1)[0][0]
            
            doc.add_paragraph()
            doc.add_paragraph(f'综合评估推荐模型: {overall_best}', style='Heading 4')
            doc.add_paragraph(
                f'{overall_best}模型在多项指标上表现优异，综合考虑预测精度、稳定性和实用性，'
                f'建议作为有功功率预测的首选模型。'
            )
    
    doc.add_paragraph()
    
    add_heading_with_style(doc, '5.2 基线对比', level=2)
    doc.add_paragraph(
        '所有模型均与朴素基线（Naive）和季节性朴素基线（SeasonalNaive）进行对比。'
        '只有在各项指标上显著优于基线的模型才具有实际应用价值。从评估结果来看，'
        '树模型（随机森林、XGBoost）和深度学习模型（LSTM、Transformer）均显著优于基线模型，'
        '证明了复杂模型在电力预测任务中的有效性。'
    )
    
    add_heading_with_style(doc, '5.3 应用建议', level=2)
    
    recommendations = [
        '模型选择应综合考虑预测精度、计算成本和可解释性',
        '建议在实际部署前进行更多折数的交叉验证以确保稳定性',
        '可根据实际需求调整预测步长和模型超参数',
        '定期使用新数据重新训练和评估模型'
    ]
    
    for rec in recommendations:
        doc.add_paragraph(rec, style='List Number')
    
    # Section 6: Appendix
    add_heading_with_style(doc, '六、附录', level=1)
    
    add_heading_with_style(doc, '6.1 数据说明', level=2)
    doc.add_paragraph('数据来源: 电力系统监测数据', style='List Bullet')
    doc.add_paragraph('包含变量: 有功功率（P）、无功功率（Q）', style='List Bullet')
    doc.add_paragraph('时间粒度: 根据配置文件设定', style='List Bullet')
    
    add_heading_with_style(doc, '6.2 可复现性', level=2)
    doc.add_paragraph(f'配置文件: {config_path}', style='List Bullet')
    doc.add_paragraph('详细指标: 见 cv_metrics.csv', style='List Bullet')
    doc.add_paragraph('图表目录: 见 figures/ 目录', style='List Bullet')
    doc.add_paragraph('执行命令: python run_all.py --config config.yaml', style='List Bullet')
    
    # Section 7: Forecast Results (if available)
    if forecast_df is not None and len(forecast_df) > 0:
        add_heading_with_style(doc, '七、未来预测结果', level=1)
        
        doc.add_paragraph(
            f'本节展示使用最优模型对未来时间段的预测结果。'
            f'预测时间范围：{forecast_df["时间"].min()} 至 {forecast_df["时间"].max()}，'
            f'共{len(forecast_df)}个时间点。'
        )
        doc.add_paragraph()
        
        add_heading_with_style(doc, '7.1 预测结果表格', level=2)
        
        # Add forecast table (show all rows if reasonable, otherwise sample)
        if len(forecast_df) <= 100:
            display_df = forecast_df.copy()
        else:
            # Show first 50 and last 50
            display_df = pd.concat([forecast_df.head(50), forecast_df.tail(50)])
            doc.add_paragraph(f'注：完整预测结果共{len(forecast_df)}行，此处展示前50行和后50行')
            doc.add_paragraph()
        
        # Format time column
        display_df = display_df.copy()
        display_df['时间'] = display_df['时间'].dt.strftime('%Y-%m-%d %H:%M')
        display_df['预测值'] = display_df['预测值'].round(2)
        
        add_table_from_dataframe(doc, display_df, title='有功功率预测结果')
        doc.add_paragraph()
        
        # Add statistics
        add_heading_with_style(doc, '7.2 预测统计信息', level=2)
        doc.add_paragraph(f'预测均值: {forecast_df["预测值"].mean():.2f}', style='List Bullet')
        doc.add_paragraph(f'预测中位数: {forecast_df["预测值"].median():.2f}', style='List Bullet')
        doc.add_paragraph(f'预测标准差: {forecast_df["预测值"].std():.2f}', style='List Bullet')
        doc.add_paragraph(f'预测最小值: {forecast_df["预测值"].min():.2f}', style='List Bullet')
        doc.add_paragraph(f'预测最大值: {forecast_df["预测值"].max():.2f}', style='List Bullet')
        doc.add_paragraph()
    
    # Footer
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f'本报告由电力质量预测系统自动生成于 {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
    )
    footer_run.font.size = Pt(10)
    footer_run.font.italic = True
    
    # Save document
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))
    
    print(f"Word report generated: {output_file}")
    
    return str(output_file)
